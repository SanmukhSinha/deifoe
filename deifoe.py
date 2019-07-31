import requests
import io
from bs4 import BeautifulSoup as bs
import docx
from docx.shared import Inches

LOGIN_URL = 'http://deifoe.org/login/index.php'  # don't change this
FILE_URL = ''
USERNAME = ''
PASSWORD = ''

USERNAME = input("Username: ")
PASSWORD = input("Password: ")
FILE_URL = input("File Address: ")
payload = {'username': USERNAME,
           'password': PASSWORD}

print("Work in Progress...")

with requests.Session() as session:
    post = session.post(LOGIN_URL, payload)
    r = session.get(FILE_URL)
    soup = bs(r.text, 'html5lib')
    data = soup.find('div', {'role': 'main'})
    name = data.h2.text

    doc = docx.Document()
    doc.add_heading(name, 0)

    p = data.find('p')
    main_tag = p.parent
    contents = main_tag.contents
    children = main_tag.children

    for tag in children:
        if tag.name == 'p':
            if tag.find('b'):
                bold = tag.find('b')
                if bold.text == tag.text:
                    heading = doc.add_heading(bold.text, 1)
                else:
                    para = None
                    for t in tag.children:
                        if t.name == "br":
                            para = doc.add_paragraph("")
                            para = None
                        else:
                            if para:
                                try:
                                    para.add_run(t.text)
                                except:
                                    para.add_run(t.string)
                            else:
                                try:
                                    para = doc.add_paragraph(t.text)
                                except:
                                    para = doc.add_paragraph(t.string)
            else:
                para = None
                for t in tag.children:
                    if t.name == "br":
                        para = doc.add_paragraph("")
                        para = None
                    else:
                        if para:
                            try:
                                para.add_run(t.text)
                            except:
                                para.add_run(t.string)
                        else:
                            try:
                                para = doc.add_paragraph(t.text)
                            except:
                                para = doc.add_paragraph(t.string)
            for img in tag.find_all('img'):
                src = img['src']
                r2 = session.get(src)
                image = io.BytesIO(r2.content)
                doc.add_picture(image, width=Inches(6))
        elif tag.name == 'img':
            src = tag['src']
            r2 = session.get(src)
            image = io.BytesIO(r2.content)
            doc.add_picture(image, width=Inches(6))
        elif tag.name == 'ol':
            for tag in tag.children:
                para = doc.add_paragraph(tag.text, style='List Bullet')
        elif tag.name == 'table':
            tbl = []
            for tr in tag.find_all('tr'):
                row = []
                for th in tr.find_all('th'):
                    row.append(th.text)
                for td in tr.find_all('td'):
                    row.append(td.text)
                tbl.append(row)

            nrow = len(tbl)
            ncol = len(tbl[0])
            table = doc.add_table(rows=nrow, cols=ncol)
            table.style='Table Grid'

            row = table.rows[0]
            i = 0
            for cell in row.cells:
                cell.text = tbl[0][i]
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                i += 1

            i = 0
            for row in table.rows:
                if i == 0:
                    i += 1
                    pass
                else:
                    j = 0
                    for cell in row.cells:
                        cell.text = tbl[i][j]
                        j += 1
                    i += 1
        else:
            para = doc.add_paragraph(tag.string)

    doc.save(name+".docx")

print("File saved as: "+name+".docx")
input("")
